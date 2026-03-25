"""Выгрузка по источнику (ac1/ac2): Word, Excel, TXT, текст для сообщения. Word требует python-docx только при выгрузке."""
import io
from typing import List

from export_excel import build_excel_bytes


def _fmt_dt(val) -> str:
    if not val:
        return ""
    s = str(val)
    return s[:19].replace("T", " ") if len(s) >= 19 else s


def build_docx_bytes(participants: List[dict], visitors: List[dict], source_label: str) -> io.BytesIO:
    try:
        from docx import Document
        from docx.shared import Pt
    except ModuleNotFoundError as e:
        raise RuntimeError(
            "Нет пакета python-docx. На сервере: pip install python-docx"
        ) from e

    doc = Document()
    doc.add_heading(f"Данные по ссылке {source_label}", level=0)
    p = doc.add_paragraph()
    p.add_run(
        "Участники (полная регистрация) и визиты без регистрации. Данные выгружены только для просмотра."
    ).font.size = Pt(10)

    doc.add_heading("Участники", level=1)
    if not participants:
        doc.add_paragraph("Нет записей.")
    else:
        table = doc.add_table(rows=1, cols=11)
        hdr = table.rows[0].cells
        headers = [
            "№",
            "Telegram ID",
            "Username",
            "Имя",
            "Возраст",
            "Деятельность",
            "Цель",
            "Телефон",
            "Регистрация",
            "Выход",
            "Источник",
        ]
        for i, h in enumerate(headers):
            hdr[i].text = h
        for row_data in participants:
            row = table.add_row().cells
            vals = [
                str(row_data.get("participant_number", "")),
                str(row_data.get("telegram_id", "")),
                str(row_data.get("username") or ""),
                str(row_data.get("name", "")),
                str(row_data.get("age", "")),
                str(row_data.get("occupation", "")),
                str(row_data.get("goal", "")),
                str(row_data.get("phone", "")),
                _fmt_dt(row_data.get("created_at")),
                _fmt_dt(row_data.get("left_at")) or "—",
                str(row_data.get("source") or ""),
            ]
            for i, v in enumerate(vals):
                row[i].text = v

    doc.add_heading("Визиты без регистрации (зашли по ссылке)", level=1)
    if not visitors:
        doc.add_paragraph("Нет записей.")
    else:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        for i, h in enumerate(["Telegram ID", "Источник", "Первый визит"]):
            hdr[i].text = h
        for row_data in visitors:
            row = table.add_row().cells
            row[0].text = str(row_data.get("telegram_id", ""))
            row[1].text = str(row_data.get("source") or "")
            row[2].text = _fmt_dt(row_data.get("first_seen_at"))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_txt_bytes(participants: List[dict], visitors: List[dict], source_label: str) -> io.BytesIO:
    lines = [
        f"Выгрузка по ссылке {source_label}",
        "Только чтение — исходные данные в базе не изменяются.",
        "",
        "=== УЧАСТНИКИ ===",
    ]
    if not participants:
        lines.append("(нет записей)")
    else:
        for p in participants:
            lines.append(
                "\t".join(
                    [
                        str(p.get("participant_number", "")),
                        str(p.get("telegram_id", "")),
                        str(p.get("username") or ""),
                        str(p.get("name", "")),
                        str(p.get("age", "")),
                        str(p.get("occupation", "")),
                        str(p.get("goal", "")),
                        str(p.get("phone", "")),
                        _fmt_dt(p.get("created_at")),
                        _fmt_dt(p.get("left_at")) or "—",
                        str(p.get("source") or ""),
                    ]
                )
            )
        lines.append("")
        lines.append("Колонки: № | Telegram ID | username | имя | возраст | деятельность | цель | телефон | регистрация | выход | источник")
    lines.extend(["", "=== ВИЗИТЫ БЕЗ РЕГИСТРАЦИИ ==="])
    if not visitors:
        lines.append("(нет записей)")
    else:
        for v in visitors:
            lines.append(
                "\t".join(
                    [
                        str(v.get("telegram_id", "")),
                        str(v.get("source") or ""),
                        _fmt_dt(v.get("first_seen_at")),
                    ]
                )
            )
        lines.append("")
        lines.append("Колонки: Telegram ID | источник | первый визит")

    text = "\n".join(lines)
    buf = io.BytesIO()
    buf.write(text.encode("utf-8-sig"))
    buf.seek(0)
    return buf


def build_message_text_chunks(
    participants: List[dict],
    visitors: List[dict],
    source_label: str,
    max_len: int = 3800,
) -> List[str]:
    """Текст для отправки в чат; разбит на части по лимиту Telegram."""
    header = (
        f"Данные по ссылке {source_label}\n\n"
        f"Участники: {len(participants)}. Визиты без регистрации: {len(visitors)}.\n"
    )
    if not participants and not visitors:
        return [header + "\nНет записей по этой ссылке."]

    lines: List[str] = [header, "", "— Участники —"]
    for p in participants:
        lines.append(
            f"№{p.get('participant_number')} | TG: {p.get('telegram_id')} | "
            f"@{p.get('username') or '—'} | {p.get('name')} | {p.get('age')} | "
            f"{p.get('occupation')} | {p.get('goal')} | {p.get('phone')} | "
            f"рег.: {_fmt_dt(p.get('created_at'))}"
        )
    lines.extend(["", "— Визиты без регистрации —"])
    if not visitors:
        lines.append("(нет записей)")
    else:
        for v in visitors:
            lines.append(f"TG: {v.get('telegram_id')} | {_fmt_dt(v.get('first_seen_at'))}")

    full = "\n".join(lines)
    chunks: List[str] = []
    pos = 0
    while pos < len(full):
        chunks.append(full[pos : pos + max_len])
        pos += max_len
    return chunks
